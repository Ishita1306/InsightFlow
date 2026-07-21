"""
Report Generator.

Compiles high-fidelity consulting-firm business reports in DOCX, HTML, and PDF (print-optimized HTML) formats.
"""

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io


def generate_svg_line_chart(forecast_df: pd.DataFrame) -> str:
    """Generate a clean vector SVG line chart representing historical actuals and forecast projections."""
    if forecast_df is None or forecast_df.empty:
        return ""
    
    # Filter valid lines
    df_clean = forecast_df.dropna(subset=["Actual", "Forecast"], how="all").copy()
    if len(df_clean) < 2:
        return ""
    
    y_min = float(df_clean[["Actual", "Forecast", "Lower Bound", "Upper Bound"]].min().min())
    y_max = float(df_clean[["Actual", "Forecast", "Lower Bound", "Upper Bound"]].max().max())
    y_range = y_max - y_min if y_max - y_min > 0 else 1.0
    
    n_points = len(df_clean)
    width = 600
    height = 200
    padding = 20
    
    hist_points = []
    fore_points = []
    upper_points = []
    lower_points = []
    
    for idx, row in df_clean.reset_index(drop=True).iterrows():
        x = padding + (idx / (n_points - 1)) * (width - 2 * padding)
        
        def map_y(val):
            if pd.isna(val):
                return None
            return height - padding - ((float(val) - y_min) / y_range) * (height - 2 * padding)
        
        act_y = map_y(row.get("Actual"))
        for_y = map_y(row.get("Forecast"))
        low_y = map_y(row.get("Lower Bound"))
        upp_y = map_y(row.get("Upper Bound"))
        
        if act_y is not None:
            hist_points.append(f"{x},{act_y}")
        if for_y is not None:
            fore_points.append(f"{x},{for_y}")
        if low_y is not None:
            lower_points.append((x, low_y))
        if upp_y is not None:
            upper_points.append((x, upp_y))
            
    hist_path = f"M {hist_points[0]} L " + " L ".join(hist_points[1:]) if hist_points else ""
    fore_path = f"M {fore_points[0]} L " + " L ".join(fore_points[1:]) if fore_points else ""
    
    conf_path = ""
    if upper_points and lower_points:
        poly_points = []
        for x, y in upper_points:
            poly_points.append(f"{x},{y}")
        for x, y in reversed(lower_points):
            poly_points.append(f"{x},{y}")
        conf_path = f"M {poly_points[0]} L " + " L ".join(poly_points[1:]) + " Z"
        
    svg_str = f"""
    <svg viewBox="0 0 {width} {height}" width="100%" height="{height}" style="background: rgba(255,255,255,0.01); border-radius: 8px; border: 1px solid rgba(255,255,255,0.05); margin-top: 10px;">
        <!-- Grid Lines -->
        <line x1="{padding}" y1="{padding}" x2="{width-padding}" y2="{padding}" stroke="rgba(255,255,255,0.04)" stroke-width="1" />
        <line x1="{padding}" y1="{height/2}" x2="{width-padding}" y2="{height/2}" stroke="rgba(255,255,255,0.04)" stroke-width="1" />
        <line x1="{padding}" y1="{height-padding}" x2="{width-padding}" y2="{height-padding}" stroke="rgba(255,255,255,0.04)" stroke-width="1" />
        
        <!-- Confidence Band -->
        {f'<path d="{conf_path}" fill="rgba(99, 102, 241, 0.08)" stroke="none" />' if conf_path else ''}
        
        <!-- Historical Line -->
        {f'<path d="{hist_path}" fill="none" stroke="#06B6D4" stroke-width="2" />' if hist_path else ''}
        
        <!-- Forecast Line -->
        {f'<path d="{fore_path}" fill="none" stroke="#6366F1" stroke-width="2" stroke-dasharray="4,4" />' if fore_path else ''}
    </svg>
    """
    return svg_str


def compile_excel_report(df: pd.DataFrame, summary: dict, audit: dict) -> bytes:
    """Generate Excel workbook bytes using pandas Excel writer or CSV fallback."""
    output = io.BytesIO()
    try:
        with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
            # Sheet 1: Executive Summary
            summary_data = {
                "Metric": ["Total Rows", "Total Columns", "Missing Cells", "Duplicate Rows", "Memory (Bytes)"],
                "Value": [summary["rows"], summary["columns"], summary["missing_cells"], summary["duplicate_rows"], summary["memory_bytes"]]
            }
            pd.DataFrame(summary_data).to_excel(writer, sheet_name="Summary", index=False)
            
            # Sheet 2: Data Quality Issues
            quality_data = {
                "Quality Check": [
                    "Missing Cells Count", "Duplicate Rows Count", "Empty Columns Count", 
                    "Incorrect Datatypes", "Outliers Detected", "Constant Columns Count",
                    "High Cardinality Columns", "Invalid Date Formats"
                ],
                "Count": [
                    audit["missing_cells"], audit["duplicate_rows"], len(audit["empty_cols"]),
                    len(audit["incorrect_cols"]), audit["outliers_count"], len(audit["constant_cols"]),
                    len(audit["high_card_cols"]), sum(audit["invalid_date_cols"].values())
                ]
            }
            pd.DataFrame(quality_data).to_excel(writer, sheet_name="Data Quality Audit", index=False)
            
            # Sheet 3: Dataset Sample
            df.head(100).to_excel(writer, sheet_name="Data Sample", index=False)
    except Exception:
        csv_str = df.head(100).to_csv(index=False)
        return csv_str.encode("utf-8")
        
    return output.getvalue()


def compile_pdf_document(filename: str, summary: dict, audit: dict, health_score: float, forecast_df: pd.DataFrame = None) -> bytes:
    """Generate a print-optimized, professional white-paper consulting HTML document representing the PDF output."""
    svg_chart = generate_svg_line_chart(forecast_df)
    
    forecast_rows = ""
    if forecast_df is not None and not forecast_df.empty:
        forecast_only = forecast_df[forecast_df["Forecast"].notna()].head(10)
        for _, r in forecast_only.iterrows():
            forecast_rows += f"""
            <tr>
                <td>{str(r["Timeline"])[:10]}</td>
                <td>{r["Forecast"]:.4f}</td>
                <td>{r["Lower Bound"]:.4f}</td>
                <td>{r["Upper Bound"]:.4f}</td>
            </tr>
            """
    else:
        forecast_rows = "<tr><td colspan='4' style='text-align: center; color: #666;'>No active forecast data computed for this session.</td></tr>"

    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>Kosvio Executive Intelligence Report - {filename}</title>
        <style>
            @media print {{
                body {{ margin: 0; background: #fff; color: #000; font-size: 11pt; }}
                .no-print {{ display: none; }}
                .page-break {{ page-break-before: always; }}
            }}
            body {{
                font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
                color: #2D3748;
                line-height: 1.6;
                margin: 40px auto;
                max-width: 800px;
                background: #FFFFFF;
                padding: 0 20px;
            }}
            .header {{
                border-bottom: 3px solid #6366F1;
                padding-bottom: 20px;
                margin-bottom: 30px;
                display: flex;
                justify-content: space-between;
                align-items: flex-end;
            }}
            .logo-area {{ text-align: left; }}
            .logo-main {{ font-size: 26px; font-weight: 800; color: #1A202C; letter-spacing: -0.03em; }}
            .logo-sub {{ font-size: 11px; text-transform: uppercase; color: #6366F1; font-weight: 700; letter-spacing: 0.1em; margin-top: 2px; }}
            .doc-meta {{ text-align: right; font-size: 12px; color: #718096; }}
            .title {{ font-size: 24px; font-weight: 700; color: #1A202C; margin-top: 20px; margin-bottom: 10px; }}
            .section {{ margin-bottom: 35px; }}
            .section-title {{
                font-size: 16px;
                font-weight: 700;
                text-transform: uppercase;
                color: #4F46E5;
                border-bottom: 1px solid #E2E8F0;
                padding-bottom: 6px;
                margin-bottom: 15px;
                letter-spacing: 0.05em;
            }}
            .grid {{
                display: grid;
                grid-template-columns: repeat(4, 1fr);
                gap: 15px;
                margin-bottom: 25px;
            }}
            .kpi-card {{
                background: #F7FAFC;
                border: 1px solid #E2E8F0;
                border-radius: 6px;
                padding: 15px;
                text-align: center;
            }}
            .kpi-val {{ font-size: 20px; font-weight: 700; color: #1A202C; margin: 4px 0 0; }}
            .kpi-lbl {{ font-size: 10px; color: #718096; text-transform: uppercase; font-weight: 600; }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin: 15px 0;
            }}
            th {{
                background: #F7FAFC;
                border-bottom: 2px solid #E2E8F0;
                font-weight: 700;
                text-align: left;
                padding: 10px;
                font-size: 12px;
                color: #4A5568;
                text-transform: uppercase;
            }}
            td {{
                padding: 10px;
                border-bottom: 1px solid #EDF2F7;
                font-size: 12px;
            }}
            .bullet {{
                margin-bottom: 10px;
                padding-left: 15px;
                text-indent: -15px;
                font-size: 13px;
            }}
            .footer {{
                margin-top: 60px;
                font-size: 11px;
                color: #A0AEC0;
                text-align: center;
                border-top: 1px solid #E2E8F0;
                padding-top: 20px;
            }}
            /* Light theme replacement for SVG inside PDF */
            svg {{
                filter: invert(0.8) hue-rotate(180deg);
                background: #F7FAFC !important;
                border: 1px solid #E2E8F0 !important;
            }}
        </style>
        <script>
            window.onload = function() {{
                // Trigger auto print dialog on load
                setTimeout(function() {{ window.print(); }}, 800);
            }}
        </script>
    </head>
    <body>
        <div class="header">
            <div class="logo-area">
                <div class="logo-main">KOSVIO</div>
                <div class="logo-sub">Strategic Advisory Group</div>
            </div>
            <div class="doc-meta">
                <strong>Confidential Report</strong><br>
                Source File: {filename}<br>
                Format: PDF Business Briefing
            </div>
        </div>

        <div class="title">Executive Data Analytics Briefing</div>
        
        <div class="section">
            <div class="section-title">1. Executive Summary</div>
            <p>This professional audit details the analytical integrity, data quality configurations, and time-series forecasting models compiled for <strong>{filename}</strong>. Our analytics ingestion pipelines have run structural validations across all recorded data metrics. The overall data asset quality has been scored at <strong>{health_score:.1f}/100</strong>, indicating model readiness.</p>
        </div>

        <div class="section">
            <div class="section-title">2. Key Performance Indicators</div>
            <div class="grid">
                <div class="kpi-card">
                    <div class="kpi-lbl">Total Samples</div>
                    <div class="kpi-val">{summary['rows']:,}</div>
                </div>
                <div class="kpi-card">
                    <div class="kpi-lbl">Measured Variables</div>
                    <div class="kpi-val">{summary['columns']}</div>
                </div>
                <div class="kpi-card">
                    <div class="kpi-lbl">Quality Health Rating</div>
                    <div class="kpi-val">{health_score:.1f}%</div>
                </div>
                <div class="kpi-card">
                    <div class="kpi-lbl">System Footprint</div>
                    <div class="kpi-val">{summary['memory_bytes'] / (1024 * 1024):.2f} MB</div>
                </div>
            </div>
        </div>

        <div class="section page-break">
            <div class="section-title">3. Data Quality Findings & Anomalies</div>
            <div class="bullet"><strong>• Redundancies & Duplicate Rows:</strong> {audit['duplicate_rows']} duplicate row sequences identified.</div>
            <div class="bullet"><strong>• Missing Values:</strong> {audit['missing_cells']} blank cells found.</div>
            <div class="bullet"><strong>• Outlier Detections (IQR):</strong> {audit['outliers_count']} extreme statistical values flagged.</div>
            <div class="bullet"><strong>• Constant Columns:</strong> {len(audit['constant_cols'])} invariant variables.</div>
        </div>

        <div class="section">
            <div class="section-title">4. Strategic Recommendations</div>
            <div class="bullet"><strong>• Imputation Rule:</strong> Impute empty numeric cells using target medians to prevent analytical skew.</div>
            <div class="bullet"><strong>• Deduplication Strategy:</strong> Clean redundant rows to stabilize variance metrics.</div>
            <div class="bullet"><strong>• Dimensional Alignment:</strong> Cast variables with mismatched types to appropriate date/categorical representations.</div>
        </div>

        {f'''
        <div class="section page-break">
            <div class="section-title">5. Time-Series Projections & Forecast</div>
            <p>The chart below displays historical trends mapped alongside 95% confidence projections computed for the forecasting horizon:</p>
            {svg_chart}
            
            <h4 style="margin-top: 25px; margin-bottom: 10px; font-size: 13px; text-transform: uppercase;">Forecast Timeline Projections</h4>
            <table>
                <thead>
                    <tr>
                        <th>Timeline</th>
                        <th>Projected Value</th>
                        <th>Lower Bound (95%)</th>
                        <th>Upper Bound (95%)</th>
                    </tr>
                </thead>
                <tbody>
                    {forecast_rows}
                </tbody>
            </table>
        </div>
        ''' if forecast_df is not None and not forecast_df.empty else ''}

        <div class="footer">
            Confidential Document &copy; 2026 Kosvio AI Business Intelligence Platform. Compiled via in-memory sessions.
        </div>
    </body>
    </html>
    """
    return html_content.encode("utf-8")


def compile_docx_report(filename: str, summary: dict, audit: dict, health_score: float, forecast_df: pd.DataFrame = None) -> bytes:
    """Generate a highly professional, beautifully formatted Word document (.docx) using python-docx."""
    doc = Document()
    
    # Configure consulting colors
    c_primary = RGBColor(79, 70, 229) # Purple
    c_dark = RGBColor(26, 32, 44)
    c_muted = RGBColor(113, 128, 150)
    
    # Title Block
    title_p = doc.add_paragraph()
    t_run = title_p.add_run("KOSVIO EXECUTIVE BRIEFING")
    t_run.font.name = "Arial"
    t_run.font.size = Pt(24)
    t_run.font.bold = True
    t_run.font.color.rgb = c_primary
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub_p = doc.add_paragraph()
    s_run = sub_p.add_run(f"Data Intelligence, Quality Audit, & Forecasting Projections for {filename}")
    s_run.font.name = "Arial"
    s_run.font.size = Pt(11)
    s_run.font.italic = True
    s_run.font.color.rgb = c_muted
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Horizontal Divider Line
    doc.add_paragraph("─" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1: Executive Summary
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Executive Summary")
    h1_run.font.color.rgb = c_primary
    h1_run.font.bold = True
    
    doc.add_paragraph(
        f"This business intelligence report outlines the analytical structures, data quality metrics, "
        f"and time-series predictions generated for the corporate data asset: {filename}. "
        f"The dataset consists of {summary['rows']:,} records and {summary['columns']} variables. "
        f"The overall data quality rating has been calculated at {health_score:.1f}% out of 100%."
    )
    
    # Section 2: KPIs Table
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("2. Key Performance Indicators")
    h2_run.font.color.rgb = c_primary
    h2_run.font.bold = True
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'KPI Metric'
    hdr_cells[1].text = 'Measured Value'
    
    kpis = [
        ("Total Rows", f"{summary['rows']:,}"),
        ("Total Columns", f"{summary['columns']}"),
        ("Quality Audit Score", f"{health_score:.1f}%"),
        ("System Footprint (MB)", f"{summary['memory_bytes'] / (1024 * 1024):.2f} MB"),
        ("Missing Cells Count", f"{audit['missing_cells']:,}"),
        ("Duplicate Row Count", f"{audit['duplicate_rows']:,}"),
        ("Outliers Detected (IQR)", f"{audit['outliers_count']:,}"),
    ]
    for metric, val in kpis:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = val
        
    doc.add_paragraph() # Spacer
    
    # Section 3: Data Quality & Recommendations
    h3 = doc.add_heading(level=1)
    h3_run = h3.add_run("3. Strategic Action Items")
    h3_run.font.color.rgb = c_primary
    h3_run.font.bold = True
    
    doc.add_paragraph("• Imputation Routine: Remediate incomplete cells using column-specific statistics to avoid downstream estimation error.")
    doc.add_paragraph("• Redundant Deduplication: Drop duplicate rows to stabilize covariance patterns.")
    doc.add_paragraph("• Dimension Corrections: Align columns possessing data type inconsistencies before running forecasts.")
    
    # Section 4: Forecast Projections
    if forecast_df is not None and not forecast_df.empty:
        h4 = doc.add_heading(level=1)
        h4_run = h4.add_run("4. Predictive Forecasting Projections")
        h4_run.font.color.rgb = c_primary
        h4_run.font.bold = True
        
        doc.add_paragraph(
            "Below are the calculated projection variables for the chosen forecasting horizon: "
        )
        
        f_table = doc.add_table(rows=1, cols=4)
        f_table.style = 'Light Shading Accent 1'
        f_hdr = f_table.rows[0].cells
        f_hdr[0].text = 'Timeline'
        f_hdr[1].text = 'Projected Value'
        f_hdr[2].text = 'Lower Bound (95%)'
        f_hdr[3].text = 'Upper Bound (95%)'
        
        forecast_only = forecast_df[forecast_df["Forecast"].notna()].head(15)
        for _, r in forecast_only.iterrows():
            row_cells = f_table.add_row().cells
            row_cells[0].text = str(r["Timeline"])[:10]
            row_cells[1].text = f"{r['Forecast']:.4f}"
            row_cells[2].text = f"{r['Lower Bound']:.4f}"
            row_cells[3].text = f"{r['Upper Bound']:.4f}"
            
    doc.add_paragraph()
    doc.add_paragraph("Confidential ─ Generated by Kosvio Enterprise AI Business Intelligence Platform.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def compile_html_report(filename: str, summary: dict, audit: dict, health_score: float, forecast_df: pd.DataFrame = None) -> bytes:
    """Generate a premium dark-themed glassmorphic consulting HTML document matching Kosvio styling."""
    svg_chart = generate_svg_line_chart(forecast_df)
    
    forecast_rows = ""
    if forecast_df is not None and not forecast_df.empty:
        forecast_only = forecast_df[forecast_df["Forecast"].notna()].head(10)
        for _, r in forecast_only.iterrows():
            forecast_rows += f"""
            <tr>
                <td>{str(r["Timeline"])[:10]}</td>
                <td style="color: #6366F1; font-weight: 600;">{r["Forecast"]:.4f}</td>
                <td>{r["Lower Bound"]:.4f}</td>
                <td>{r["Upper Bound"]:.4f}</td>
            </tr>
            """
    else:
        forecast_rows = "<tr><td colspan='4' style='text-align: center; color: var(--subtext);'>No active forecast data computed for this session. Run a forecast in the workspace first.</td></tr>"

    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>Kosvio Executive Briefing - {filename}</title>
        <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap" rel="stylesheet">
        <style>
            :root {{
                --bg: #030014;
                --card: rgba(255, 255, 255, 0.03);
                --border: rgba(255, 255, 255, 0.08);
                --text: #F3F4F6;
                --subtext: #9CA3AF;
                --primary: #6366F1;
                --accent: #06B6D4;
            }}
            body {{
                background-color: var(--bg);
                color: var(--text);
                font-family: 'Inter', sans-serif;
                margin: 0;
                padding: 40px 20px;
                display: flex;
                justify-content: center;
            }}
            .report-container {{
                max-width: 850px;
                width: 100%;
            }}
            .glass-card {{
                background: var(--card);
                border: 1px solid var(--border);
                border-radius: 16px;
                padding: 30px;
                backdrop-filter: blur(16px);
                -webkit-backdrop-filter: blur(16px);
                box-shadow: 0 8px 32px 0 rgba(0, 0, 0, 0.4);
                margin-bottom: 25px;
            }}
            .header {{
                display: flex;
                justify-content: space-between;
                align-items: center;
                border-bottom: 1px solid var(--border);
                padding-bottom: 20px;
                margin-bottom: 30px;
            }}
            .logo-main {{
                font-size: 28px;
                font-weight: 800;
                color: #FFFFFF;
                letter-spacing: -0.03em;
                background: linear-gradient(135deg, #6366F1, #06B6D4);
                -webkit-background-clip: text;
                -webkit-text-fill-color: transparent;
            }}
            .logo-sub {{
                font-size: 10px;
                text-transform: uppercase;
                color: var(--accent);
                font-weight: 700;
                letter-spacing: 0.15em;
                margin-top: 4px;
            }}
            .doc-meta {{
                text-align: right;
                font-size: 12px;
                color: var(--subtext);
                line-height: 1.5;
            }}
            .title {{
                font-size: 26px;
                font-weight: 800;
                color: #FFFFFF;
                margin-bottom: 10px;
            }}
            .subtitle {{
                font-size: 14px;
                color: var(--subtext);
                margin-bottom: 30px;
            }}
            .section-title {{
                font-size: 15px;
                font-weight: 700;
                text-transform: uppercase;
                color: var(--primary);
                letter-spacing: 0.08em;
                margin-bottom: 15px;
                border-left: 3px solid var(--primary);
                padding-left: 10px;
            }}
            .grid {{
                display: grid;
                grid-template-columns: repeat(4, 1fr);
                gap: 15px;
                margin-bottom: 20px;
            }}
            .kpi-card {{
                background: rgba(255, 255, 255, 0.015);
                border: 1px solid var(--border);
                border-radius: 12px;
                padding: 15px;
                text-align: center;
                transition: transform 0.2s;
            }}
            .kpi-card:hover {{
                transform: translateY(-2px);
                border-color: var(--primary);
            }}
            .kpi-val {{ font-size: 20px; font-weight: 800; color: #FFFFFF; margin: 4px 0 0; }}
            .kpi-lbl {{ font-size: 9px; color: var(--subtext); text-transform: uppercase; font-weight: 600; letter-spacing: 0.05em; }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin: 20px 0;
            }}
            th {{
                background: rgba(255, 255, 255, 0.02);
                border-bottom: 2px solid var(--border);
                font-weight: 700;
                text-align: left;
                padding: 12px;
                font-size: 11px;
                color: var(--text);
                text-transform: uppercase;
                letter-spacing: 0.05em;
            }}
            td {{
                padding: 12px;
                border-bottom: 1px solid var(--border);
                font-size: 12px;
                color: var(--subtext);
            }}
            tr:hover td {{
                color: #FFFFFF;
                background: rgba(255, 255, 255, 0.01);
            }}
            .bullet {{
                margin-bottom: 12px;
                font-size: 13px;
                color: var(--subtext);
                line-height: 1.5;
            }}
            .bullet strong {{
                color: #FFFFFF;
            }}
            .footer {{
                margin-top: 50px;
                font-size: 11px;
                color: rgba(255, 255, 255, 0.2);
                text-align: center;
                border-top: 1px solid var(--border);
                padding-top: 20px;
            }}
        </style>
    </head>
    <body>
        <div class="report-container">
            <div class="glass-card">
                <div class="header">
                    <div>
                        <div class="logo-main">KOSVIO</div>
                        <div class="logo-sub">AI Business Intelligence Platform</div>
                    </div>
                    <div class="doc-meta">
                        <strong>Confidential Report</strong><br>
                        Asset: {filename}<br>
                        Format: Premium HTML Studio
                    </div>
                </div>

                <div class="title">Strategic Briefing & Forecasting Analytics</div>
                <div class="subtitle">An automated diagnostic audit of data configurations, anomalies, and prediction models.</div>

                <div class="section-title">1. Executive Summary</div>
                <p style="font-size: 14px; line-height: 1.6; color: var(--subtext); margin-bottom: 30px;">
                    This comprehensive analysis validates the integrity, quality metrics, and future trends of <strong>{filename}</strong>. 
                    The ingestion layer validated formatting variables, empty records, and datatype alignments. 
                    The final data quality rating of this asset is <strong>{health_score:.1f}/100</strong>, indicating that the data is ready for predictive modeling.
                </p>

                <div class="section-title">2. Ingested Analytics KPIs</div>
                <div class="grid">
                    <div class="kpi-card">
                        <div class="kpi-lbl">Total Samples</div>
                        <div class="kpi-val">{summary['rows']:,}</div>
                    </div>
                    <div class="kpi-card">
                        <div class="kpi-lbl">Variables</div>
                        <div class="kpi-val">{summary['columns']}</div>
                    </div>
                    <div class="kpi-card">
                        <div class="kpi-lbl">Quality Score</div>
                        <div class="kpi-val" style="color: var(--accent);">{health_score:.1f}%</div>
                    </div>
                    <div class="kpi-card">
                        <div class="kpi-lbl">Footprint</div>
                        <div class="kpi-val">{summary['memory_bytes'] / (1024 * 1024):.2f} MB</div>
                    </div>
                </div>
            </div>

            <div class="glass-card">
                <div class="section-title">3. Data Quality Findings & Actions</div>
                <div class="bullet"><strong>• Redundancies & Duplicates:</strong> {audit['duplicate_rows']} duplicate row sequences identified. Clean rows to stabilize variance.</div>
                <div class="bullet"><strong>• Missing Values:</strong> {audit['missing_cells']} empty cells found. Impute using median statistics to avoid bias.</div>
                <div class="bullet"><strong>• Statistical Outliers:</strong> {audit['outliers_count']} extreme outliers mapped using standard IQR methods. Flag outliers to stabilize forecasts.</div>
                <div class="bullet"><strong>• Datatype Inconsistencies:</strong> {len(audit['incorrect_cols'])} columns contain datatype inconsistencies. Align formats for accurate timelines.</div>
            </div>

            {f'''
            <div class="glass-card">
                <div class="section-title">4. Time-Series Projections & Predictions</div>
                <p style="font-size: 13px; color: var(--subtext); margin-bottom: 20px;">The interactive vector chart below tracks historical actuals alongside the 95% confidence forecasting projections:</p>
                
                {svg_chart}

                <h4 style="margin-top: 30px; margin-bottom: 10px; font-size: 12px; text-transform: uppercase; color: var(--text); letter-spacing: 0.05em;">Timeline Horizon Forecast</h4>
                <table>
                    <thead>
                        <tr>
                            <th>Timeline</th>
                            <th>Projected Value</th>
                            <th>Lower Bound (95%)</th>
                            <th>Upper Bound (95%)</th>
                        </tr>
                    </thead>
                    <tbody>
                        {forecast_rows}
                    </tbody>
                </table>
            </div>
            ''' if forecast_df is not None and not forecast_df.empty else ''}

            <div class="footer">
                Kosvio AI Business Intelligence Platform &copy; 2026. All rights reserved. Confidential.
            </div>
        </div>
    </body>
    </html>
    """
    return html_content.encode("utf-8")


def compile_powerpoint_briefing(filename: str, summary: dict, audit: dict, health_score: float) -> bytes:
    """Generate a clean slides outline representation representing the PowerPoint briefing."""
    briefing = f"""# Kosvio EXECUTIVE BRIEFING DECK: {filename}
# Slide 1: Title Slide
- Title: Kosvio Data Intelligence Briefing
- Subtitle: Structural Analysis & Quality Profile for {filename}
- Date: Compiled Workspace Session

# Slide 2: Structural Dimensions
- Title: Ingested Dataset Overview
- Bullet 1: Total Processed Samples: {summary['rows']:,} rows
- Bullet 3: Memory Footprint: {summary['memory_bytes'] / 1024:.1f} KB
- Bullet 4: Active RAM footprint: {100.0 - summary['missing_pct']:.2f}% completeness

# Slide 3: Data Quality Audit
- Title: Anomalies & Data Integrity Check
- Health Rating: {health_score:.1f}/100 Health Score
- Bullet 1: Missing values count: {audit['missing_cells']:,} cells
- Bullet 2: Redundant rows count: {audit['duplicate_rows']:,} duplicates
- Bullet 3: Outliers detected: {audit['outliers_count']:,} values outside IQR
- Bullet 4: Zero-entropy constant variables: {len(audit['constant_cols'])} columns

# Slide 4: Actionable Recommendations
- Title: Strategic Steps & Imputations
- Recommendation 1: Perform deduplication and remove duplicate records.
- Recommendation 2: Trim leading and trailing spaces from object columns.
- Recommendation 3: Impute missing numerical cells with median values.
- Recommendation 4: Standardize dates and datatype alignments.
"""
    return briefing.encode("utf-8")
